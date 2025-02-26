import traceback
import streamlit as st
import json
import datetime
import time
import re
import io
from streamlit_float import float_init
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from services.time import date_to_chinese
from services.tools import replace_card_numbers
from services.ds_service import DSService
from services.chat_service import ChatService
from config import RESOURCES_DIR

def use_re(text):
    """从文本中提取 JSON 数据"""
    json_pattern = re.compile(r'\{.*\}', re.DOTALL)
    json_match = json_pattern.search(text)
    if json_match:
        try:
            return json.loads(json_match.group())
        except json.JSONDecodeError as e:
            st.error(f"JSON 解码错误: {e}")
            return None
    return None

def process_data_list(data_list, table, columns):
    """处理数据并填充表格"""
    try:
        if "Preservation_Results" not in data_list:
            raise KeyError("'Preservation_Results' 不在 data_list 中")

        for item in data_list["Preservation_Results"]:
            if not isinstance(item, dict):
                print(f"警告: item 应该是一个字典，但实际是 {type(item)}")
                continue

            row_cells = table.add_row().cells
            for i, column in enumerate(columns):
                if column not in item:
                    print(f"警告: 列 {column} 不在 item 中: {item}")
                    row_cells[i].text = str(item.get("Case,No", 'N/A'))
                else:
                    row_cells[i].text = str(item[column])
                
                if column == "Preservation_Measure":
                    row_cells[i].text = replace_card_numbers(row_cells[i].text)
                
                # 设置单元格格式
                row_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = row_cells[i].paragraphs[0].runs[0]
                run.font.name = "仿宋_GB2312"
                run.font.size = Pt(10.5)
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "仿宋_GB2312")
        return table
    except Exception as e:
        print(f"发生错误: {e}")
        traceback.print_exc()

def replace_in_paragraph(doc, data):
    """替换文档段落中的占位符，并保留格式"""
    for paragraph in doc.paragraphs:
        for key, value in data.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

def write_file(data_list):
    """根据传入的数据生成并保存 Word 文档"""
    doc = Document(Path(RESOURCES_DIR) / "template.docx")
    data = {
        "{Applicant}": data_list["Applicant"],
        "{Respondent}": ", ".join(data_list["Respondent"]) if isinstance(data_list["Respondent"], list) else data_list["Respondent"],
        "{CBN}": data_list["Civil_Ruling_Number"],
        "{today}": date_to_chinese(datetime.date.today()),
    }
    replace_in_paragraph(doc, data)
    process_data_list(data_list, doc.tables[0], data_list["Preservation_Results"][0].keys())
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def fetch_streaming_response(completion):
    """模拟流式响应"""
    for chunk in completion:
        if chunk.choices and chunk.choices[0].delta.content:
            yield chunk.choices[0].delta.content
            time.sleep(0.25)  # 控制显示速度

def initialize_session_state():
    """初始化会话状态"""
    if "messages" not in st.session_state:
        st.session_state.messages = [
            {"role": "system", "content": """### 角色设定\n最高法院认证的司法文书结构化处理专家，具备民商事保全裁定自动化解析资质。\n\n### 任务\n 根据用户提供的文书内容，自动解析并生成结构化、标准化JSON数据输出。请直接输出最终答案，无需包含分析步骤或推理过程。输出格式需严格遵循 [指定格式]，仅呈现结论性内容。\n\n### 任务分解\n 1. **字段提取**\n    - 申请人：定位"申请人"字段，若无明确标识则赋值为"未知"\n    - 被申请人：强制要求从文书中直接提取，禁止推断或默认值\n     - 民事裁定书编号：识别"民初"、"民终"等标准编号格式\n     - 保全条目：提取所有独立保全项（包括执行失败的情况），包含：\n         * 序号\n        * 权利人(权利人为被申请人）\n        * 保全措施（"结构化模板": 操作动词[冻结/查封]+[被申请人]+标的物[金融账户/房产/动产]+实际处理结果[账户XX元]"；"多标的分割符": ["；", "。", "，"]，对多标的分开列目）\n        * 实施时间（格式：YYYY-MM-DD，无法确认时默认1970-01-01，禁止推断，如果存在明确的到期时间和期限年份，可以根据计算逻辑：到期日-期限年份 得到实施时间）\n        * 保全期限（单位：年，示例格式：一年）\n        * 到期日（计算逻辑：实施时间+期限年份）\n        * 备注（特定格式：支付宝ID/微信订单号，和执行失败原因说明）\n 2. **数据处理规范** \n    - 金融账户处理：\n        * 实际处理结果:删除应冻结内容和实际冻结字样，只能用冻结作为操作动词，用实际冻结金额作为处理结果\n        * 银行账户：只输出银行卡号，不关联归属银行信息\n        * 支付宝账户：使用机构名全称“支付宝（中国）网络技术有限公司”，加“支付宝账户”\n        * 财付通账户：使用全称”财付通支付科技有限公司“，加微信号，删除财付通账号\n         * 超额冻结处理：计算有效金额作为实际冻结金额\n    - 房产：查封房产需关联不动产登记编号\n    - 支付宝账户备注：强制列出"ID:"+识别值。"匹配'ID:'后接8-24位数字的典型格式"，示例为"ID:44329891"\n    - 微信/财付通账户备注："冻结单号:"+识别值\n    - 执行失败的情况：备注中说明具体原因\n 3. **校验与修正**\n    - 完整性检查：对提取到的信息进行校验，确保信息完整，同时列出执行失败[查封失败/冻结失败]的保全结果\n    - 时间校验：自动修正不符合YYYY-MM-DD格式的时间字段\n    - 逻辑验证：确保保全到期日=实施时间+保全期限\n    - 备注验证： 禁止添加‘余额不足’备注\n\n ### 输出规范 \n```json\n {\n    "Applicant": "张三", \n    "Respondent": "李四",\n    "Civil_Ruling_Number": "(2024)京0105民初1234号",\n    "Preservation_Results":[\n        {\n            "Case_No": 1,\n            "Right_Holder": "李四",\n            "Preservation_Measure": "冻结被申请人李四名下的支付宝（中国）网络技术有限公司1228349357支付宝账户内存款1100元",\n            "Implementation_Date": "2024-02-01",\n            "Duration_Years": "一年",\n            "Expiration_Date": "2025-01-31",\n            "Remarks": "ID:13256332" \n        }\n    ]\n} """},
        ]
    if "chat_service" not in st.session_state:
        st.session_state.chat_service = ChatService()
    if "executed" not in st.session_state:
        st.session_state.executed = True
    if "last_json_list" not in st.session_state:
        st.session_state.last_json_list = None

def initialize_dssession_state():
    if "dsmessages" not in st.session_state:
        st.session_state.dsmessages = [
            {"role": "system", "content": "你是一个AI助手，你的任务是回答用户的问题。"}
        ]
    if "ds_serivice" not in st.session_state:
        st.session_state.ds_service = DSService()

def render_chat2_interface():
    """渲染聊天界面"""
    float_init()
    st.subheader("消息区")
    initialize_dssession_state()

    # 消息显示区域
    messages_container = st.container()
    with messages_container:
        expander = st.expander("消息记录", expanded=False)
        with expander:
            for message in st.session_state.dsmessages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

    # 用户输入区域
    input_container = st.container()
    with input_container:
        if prompt := st.chat_input("输入问题"):
            with messages_container:
                # 添加用户消息
                with st.chat_message("user"):
                    st.markdown(prompt)
                st.session_state.dsmessages.append({"role": "user", "content": prompt})

                # 获取助手回复
                try:
                    start_time = time.time()
                    
                    reasoning_content = ""
                    content = ""
                    with messages_container:
                        expander_thinking = st.expander("深度思考", expanded=False)
                        with expander_thinking:
                            reasoning_placeholder = st.empty()
                        with st.chat_message("assistant"):
                            response_placeholder = st.empty()
                            for chunk in st.session_state.ds_service.get_response(st.session_state.dsmessages):
                                if hasattr(chunk.choices[0].delta, 'reasoning_content') and chunk.choices[0].delta.reasoning_content:
                                    reasoning_content += chunk.choices[0].delta.reasoning_content
                                    reasoning_placeholder.markdown(reasoning_content)
                                else:          
                                    content += chunk.choices[0].delta.content
                                    response_placeholder.markdown(content)

                    end_time = time.time()
                    elapsed_time = end_time - start_time
                    print(f"获取回答耗时: {elapsed_time:.2f} 秒")
                    st.session_state.dsmessages.append({"role": "assistant", "content": content})

                except Exception as e:
                    st.error(f"获取回答失败: {str(e)}")

    # 将输入容器固定在底部
    input_container.container().float("bottom: 0")

def render_chat_interface():
    """渲染聊天界面"""
    float_init()
    st.subheader("消息区")
    initialize_session_state()

    # 消息显示区域
    messages_container = st.container()
    with messages_container:
        expander = st.expander("消息记录", expanded=False)
        with expander:
            for message in st.session_state.messages:
                with st.chat_message(message["role"]):
                    st.markdown(message["content"])

    # 用户输入区域
    input_container = st.container()
    with input_container:
        if prompt := st.chat_input("输入问题"):
            with messages_container:
                # 处理文件 ID
                if "file_id" in st.session_state and st.session_state.file_id != 0 and st.session_state.executed:
                    st.session_state.messages.extend([
                        {"role": "system", "content": f'fileid://{st.session_state.file_id}'},
                        {"role": "user", "content": "对新上传的文件进行处理，输出结果"},
                    ])
                    st.session_state.executed = False
                else:
                    # 添加用户消息
                    with st.chat_message("user"):
                        st.markdown(prompt)
                    st.session_state.messages.append({"role": "user", "content": prompt})

                # 获取助手回复
                try:
                    start_time = time.time()
                    with messages_container:
                        with st.chat_message("system"):
                            response_placeholder = st.empty()
                            response = ""

                            for chunk in st.session_state.chat_service.get_response(st.session_state.messages):
                                response += chunk.choices[0].delta.content
                                response_placeholder.markdown(response)

                    end_time = time.time()
                    elapsed_time = end_time - start_time
                    print(f"获取回答耗时: {elapsed_time:.2f} 秒")
                    st.session_state.messages.append({"role": "system", "content": response})

                    # 解析 JSON 并生成文件
                    json_list = use_re(response)
                    if json_list:
                        st.session_state.file_stream = write_file(json_list)
                        if json_list != st.session_state.last_json_list and st.session_state.last_json_list is not None:
                            st.session_state.last_json_list = json_list
                            st.success("检测到 JSON 数据更新，已重新生成文件！")
                        else:
                            st.success("已生成文件！")

                except Exception as e:
                    st.error(f"获取回答失败: {str(e)}")

    # 将输入容器固定在底部
    input_container.container().float("bottom: 0")